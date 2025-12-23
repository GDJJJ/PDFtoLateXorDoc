import os
import cv2
import numpy as np
from pathlib import Path
from paddleocr import PPStructure, PaddleOCR
import requests
import re


def clean_text_for_xml(text):
    """移除 XML 不兼容字符"""
    if not text:
        return ""
    # 移除空字节和控制字符
    text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F]', '', text)
    return text


def process_text_with_deepseek(ocr_text):
    """使用 DeepSeek API 处理 OCR 文本，生成 LaTeX 格式"""
    url = "https://api.deepseek.com/v1/chat/completions"
    api_key = "sk-be0d106e6aac41cf8aa45299a3af7043"

    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }

    payload = {
        "model": "deepseek-chat",
        "messages": [{
            "role": "user",
            "content": f"""请对以下OCR识别文本进行处理，要求：
1. 去除识别文档的眉页(header)和眉尾(footer)内容
2. 对正文进行「最小化修正」，用LaTeX格式标记标题，仅修改明显错误的词汇（形近字、语义矛盾词）,保留正确的缩进和换行
3. 如果发现表格内容，请使用LaTeX表格语法（\\begin{{table}}...\\end{{table}}）表示
4. 具体修正要求：
   - 只修改确认错误的词，修正后的词必须与原词外形相似（如"木未"→"木末"），具体修改情况不用在生成的文本里说明
   - 不删除、不新增任何内容，不调整段落或标点
   - 去掉不必要的空格和换行
   - 移除所有 XML/LaTeX 不兼容字符，如空字节（\\x00）和控制字符
   - 去除页头页尾
   - 特殊字符需要转义：& → \\&, % → \\%, $ → \\$, # → \\#, ^ → \\^{{}}, _ → \\_, {{ → \\{{, }} → \\}}, \\ → \\textbackslash
5. 标注标题（需有明显标识，如"一、"、"二、"、"（一）、"），标题一般不是一个句子
   - 一级标题使用 \\section{{标题内容}}
   - 二级标题使用 \\subsection{{标题内容}}
   - 三级标题使用 \\subsubsection{{标题内容}}
6. 段落之间使用空行分隔，每段开头不需要特殊空格
7. 不要自行添加额外的说明（如这是修改结果或者具体修正的情况）
8. 最终格式必须是纯LaTeX代码，不要包含markdown格式，示例：
\\section{{文档标题}}
正文内容段落...

\\begin{{table}}[h]
\\centering
\\begin{{tabular}}{{cc}}
\\hline
项目名称 & 技术领域 \\\\
\\hline
馈能式智能悬架 & 先进制造 \\\\
\\hline
\\end{{tabular}}
\\caption{{表格标题}}
\\end{{table}}

待处理文本：
{ocr_text}"""
        }],
        "max_tokens": 4000,
        "temperature": 0.3
    }

    try:
        response = requests.post(url, headers=headers, json=payload, timeout=60)
        if response.status_code == 200:
            return response.json()["choices"][0]["message"]["content"]
        else:
            print(f"DeepSeek API 调用失败: {response.status_code}, {response.text}")
            return ocr_text  # 失败时返回原始文本
    except Exception as e:
        print(f"DeepSeek API 调用异常: {e}")
        return ocr_text  # 异常时返回原始文本


def save_region_image(img_array, bbox, region_type, output_dir, page_index, region_index):
    """保存区域图片"""
    try:
        if img_array is not None and len(img_array.shape) == 3:
            x1, y1, x2, y2 = map(int, bbox)
            x1, y1 = max(0, x1), max(0, y1)
            x2, y2 = min(img_array.shape[1], x2), min(img_array.shape[0], y2)
            
            if x2 > x1 and y2 > y1:
                cropped = img_array[y1:y2, x1:x2]
                # 创建类型目录
                type_dir = output_dir / region_type
                type_dir.mkdir(exist_ok=True)
                
                img_path = type_dir / f"page_{page_index}_region_{region_index}.jpg"
                cv2.imwrite(str(img_path), cropped)
                # 返回相对路径，使用正斜杠（Markdown 兼容）
                rel_path = img_path.relative_to(output_dir)
                return str(rel_path).replace('\\', '/')
    except Exception as e:
        print(f"保存区域图片失败: {e}")
    return None


def process_single_image(image_path, output_dir, page_index, structure_model, ocr_model):
    """处理单张图片"""
    print(f"处理图片: {image_path}")
    
    image_path = Path(image_path)
    
    # 读取图片
    img_array = cv2.imread(str(image_path))
    if img_array is None:
        print(f"无法读取图片: {image_path}")
        return ""
    
    # 使用 PaddleStructure 进行版面检测
    try:
        structure_results = structure_model(str(image_path))
    except Exception as e:
        print(f"版面检测失败: {e}")
        return ""
    
    latex_parts = []
    region_index = 0
    
    # 处理结果可能是列表或嵌套结构
    if not isinstance(structure_results, list):
        structure_results = [structure_results]
    
    # 展平结果并提取有效数据
    processed_results = []
    for res in structure_results:
        if isinstance(res, dict):
            processed_results.append(res)
        elif hasattr(res, '__dict__'):
            # 如果是对象，尝试转换为字典
            try:
                res_dict = {
                    'type': getattr(res, 'type', 'unknown'),
                    'bbox': getattr(res, 'bbox', []),
                    'res': getattr(res, 'res', None)
                }
                processed_results.append(res_dict)
            except:
                continue
    
    # 按位置排序（从上到下，从左到右）
    sorted_results = sorted(processed_results, key=lambda x: (
        x.get('bbox', [0, 0, 0, 0])[1] if isinstance(x, dict) and 'bbox' in x and len(x.get('bbox', [])) >= 2 else 0,
        x.get('bbox', [0, 0, 0, 0])[0] if isinstance(x, dict) and 'bbox' in x and len(x.get('bbox', [])) >= 1 else 0
    ))
    
    for res in sorted_results:
        if not isinstance(res, dict):
            continue
            
        region_type = res.get('type', 'unknown')
        bbox = res.get('bbox', [])
        
        if not bbox or len(bbox) < 4:
            continue
        
        # 文本区域：进行 OCR 并调用 DeepSeek API
        if region_type in ['text', 'title']:
            # 提取文本区域进行 OCR
            ocr_text = ""
            
            # 尝试从结果中提取 OCR 文本
            if 'res' in res and res['res']:
                ocr_text_parts = []
                for line_result in res['res']:
                    if isinstance(line_result, (list, tuple)) and len(line_result) >= 2:
                        if isinstance(line_result[1], (list, tuple)) and len(line_result[1]) >= 1:
                            text = line_result[1][0]
                        else:
                            text = str(line_result[1])
                        ocr_text_parts.append(text)
                ocr_text = "\n".join(ocr_text_parts)
            
            # 如果没有 OCR 结果或结果为空，手动进行 OCR
            if not ocr_text.strip():
                x1, y1, x2, y2 = map(int, bbox)
                x1, y1 = max(0, x1), max(0, y1)
                x2, y2 = min(img_array.shape[1], x2), min(img_array.shape[0], y2)
                
                if x2 > x1 and y2 > y1:
                    cropped = img_array[y1:y2, x1:x2]
                    try:
                        ocr_result = ocr_model.ocr(cropped, cls=True)
                        ocr_text_parts = []
                        if ocr_result and ocr_result[0]:
                            for line in ocr_result[0]:
                                if len(line) >= 2:
                                    if isinstance(line[1], (list, tuple)) and len(line[1]) >= 1:
                                        ocr_text_parts.append(line[1][0])
                                    else:
                                        ocr_text_parts.append(str(line[1]))
                        ocr_text = "\n".join(ocr_text_parts)
                    except Exception as e:
                        print(f"OCR 识别失败: {e}")
                        ocr_text = ""
            
            if ocr_text.strip():
                # 清理文本
                ocr_text = clean_text_for_xml(ocr_text)
                # 调用 DeepSeek API 处理
                print(f"  处理文本区域 {region_index}，文本长度: {len(ocr_text)}")
                processed_text = process_text_with_deepseek(ocr_text)
                latex_parts.append(processed_text)
        
        # 表格区域：先保存图片，然后尝试 OCR 表格内容
        elif region_type == 'table':
            saved_path = save_region_image(
                img_array, bbox, region_type, output_dir, page_index, region_index
            )
            if saved_path:
                # 使用 LaTeX 格式引用表格图片
                # 将路径中的反斜杠转换为正斜杠，LaTeX 使用正斜杠
                latex_path = saved_path.replace('\\', '/')
                latex_parts.append(f"\n\\begin{{figure}}[h]\n\\centering\n\\includegraphics[width=0.8\\textwidth]{{{latex_path}}}\n\\caption{{表格}}\n\\end{{figure}}\n")
            region_index += 1
        
        # 其他非文本区域：保存图片并在 LaTeX 中引用
        elif region_type in ['figure', 'image', 'formula']:
            saved_path = save_region_image(
                img_array, bbox, region_type, output_dir, page_index, region_index
            )
            if saved_path:
                # 使用 LaTeX 格式引用图片
                latex_path = saved_path.replace('\\', '/')
                caption_map = {
                    'figure': '图',
                    'image': '图片',
                    'formula': '公式'
                }
                caption = caption_map.get(region_type, '图片')
                latex_parts.append(f"\n\\begin{{figure}}[h]\n\\centering\n\\includegraphics[width=0.8\\textwidth]{{{latex_path}}}\n\\caption{{{caption}}}\n\\end{{figure}}\n")
            region_index += 1
    
    return "\n\n".join(latex_parts)


def process_all_images(input_folder, output_path, output_format='html'):
    """处理文件夹中的所有图片
    
    Args:
        input_folder: 输入图片文件夹路径
        output_path: 输出文件路径
        output_format: 输出格式，'html', 'docx', 'pdf', 或 'tex'
    """
    input_folder = Path(input_folder)
    output_path = Path(output_path)
    output_dir = output_path.parent / "processed_regions"
    output_dir.mkdir(exist_ok=True)
    
    # 初始化模型
    print("初始化 PaddleStructure 模型...")
    structure_model = PPStructure(
        layout=True,
        table=True,
        ocr=True,
        lang="ch",  # 中文文档
        structure_version='PP-StructureV2'
    )
    
    print("初始化 PaddleOCR 模型...")
    ocr_model = PaddleOCR(use_angle_cls=True, lang="ch")
    
    # 获取所有图片文件
    image_extensions = {'.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif'}
    image_files = sorted([
        f for f in input_folder.iterdir() 
        if f.suffix.lower() in image_extensions
    ])
    
    if not image_files:
        print(f"在 {input_folder} 中未找到图片文件")
        return
    
    print(f"找到 {len(image_files)} 张图片")
    
    all_latex_content = []
    
    # 处理每张图片
    for page_index, image_file in enumerate(image_files, 1):
        print(f"\n处理第 {page_index}/{len(image_files)} 张图片...")
        latex_content = process_single_image(
            image_file, output_dir, page_index, structure_model, ocr_model
        )
        if latex_content.strip():
            all_latex_content.append(f"\\subsection{{第 {page_index} 页}}\n\n{latex_content}")
    
    # 合并所有内容
    body_content = "\n\n\\newpage\n\n".join(all_latex_content)
    
    if output_format == 'html':
        # 生成 HTML 文档
        try:
            html_parts = ['<!DOCTYPE html>\n<html lang="zh-CN">\n<head>\n<meta charset="UTF-8">\n<title>文档处理结果</title>\n<style>body{font-family:"Microsoft YaHei",Arial;max-width:900px;margin:0 auto;padding:20px;line-height:1.6}h1{color:#2c3e50;border-bottom:2px solid #3498db;padding-bottom:10px}h2{color:#34495e;margin-top:30px}img{max-width:100%;height:auto;margin:20px 0;border:1px solid #ddd}p{margin:10px 0;text-indent:2em}</style>\n</head>\n<body>\n<h1>文档处理结果</h1>\n']
            
            for page_content in all_latex_content:
                page_match = re.search(r'\\subsection\{第 (\d+) 页\}', page_content)
                if page_match:
                    html_parts.append(f'<h2>第 {page_match.group(1)} 页</h2>\n')
                    page_content = re.sub(r'\\subsection\{第 \d+ 页\}\s*\n\s*', '', page_content)
                
                content_blocks = re.split(r'(\\begin\{figure\}.*?\\end\{figure\}|\\section\{[^}]+\}|\\subsection\{[^}]+\})', page_content, flags=re.DOTALL)
                
                for block in content_blocks:
                    if not block.strip():
                        continue
                    
                    if '\\begin{figure}' in block:
                        img_match = re.search(r'\\includegraphics.*?\{([^}]+)\}', block)
                        if img_match:
                            img_rel_path = img_match.group(1).replace('/', os.sep)
                            img_path = output_dir / img_rel_path
                            if img_path.exists():
                                import shutil
                                img_output_path = output_path.parent / img_rel_path.replace(os.sep, '_')
                                shutil.copy2(img_path, img_output_path)
                                html_parts.append(f'<img src="{img_output_path.name}" alt="图片">\n')
                        caption_match = re.search(r'\\caption\{([^}]+)\}', block)
                        if caption_match:
                            html_parts.append(f'<p style="text-align:center;font-style:italic;color:#666">{caption_match.group(1)}</p>\n')
                        continue
                    
                    if block.startswith('\\section{'):
                        title_text = re.sub(r'\\section\{([^}]+)\}', r'\1', block)
                        html_parts.append(f'<h2>{title_text.strip()}</h2>\n')
                    elif block.startswith('\\subsection{'):
                        title_text = re.sub(r'\\subsection\{([^}]+)\}', r'\1', block)
                        html_parts.append(f'<h3>{title_text.strip()}</h3>\n')
                    else:
                        lines = block.split('\n')
                        for line in lines:
                            line = line.strip()
                            if not line or (line.startswith('\\') and not line.startswith('\\text')):
                                continue
                            text = re.sub(r'\\textbf\{([^}]+)\}', r'<strong>\1</strong>', line)
                            text = re.sub(r'\\textit\{([^}]+)\}', r'<em>\1</em>', text)
                            text = re.sub(r'\\([&%$#_{}])', r'\1', text)
                            text = re.sub(r'\\textbackslash', '\\', text)
                            text = re.sub(r'\\[a-zA-Z]+\{([^}]+)\}', r'\1', text)
                            text = re.sub(r'\\[a-zA-Z]+', '', text)
                            if text.strip():
                                html_parts.append(f'<p>{text.strip()}</p>\n')
            
            html_parts.append('</body>\n</html>')
            html_content = ''.join(html_parts)
            html_path = output_path.with_suffix('.html')
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            print(f"\nHTML 文件已保存: {html_path}")
            return html_path
        except Exception as e:
            import traceback
            print(f"生成 HTML 失败: {e}\n{traceback.format_exc()}")
            output_format = 'tex'
    
    elif output_format == 'pdf':
        # PDF 生成：先生成 Word，用户可手动转换
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            temp_docx = output_path.with_suffix('.docx')
            doc = Document()
            doc.add_heading('文档处理结果', 0)
            
            for page_content in all_latex_content:
                page_match = re.search(r'\\subsection\{第 (\d+) 页\}', page_content)
                if page_match:
                    doc.add_heading(f'第 {page_match.group(1)} 页', level=2)
                    page_content = re.sub(r'\\subsection\{第 \d+ 页\}\s*\n\s*', '', page_content)
                
                content_blocks = re.split(r'(\\begin\{figure\}.*?\\end\{figure\}|\\section\{[^}]+\}|\\subsection\{[^}]+\})', page_content, flags=re.DOTALL)
                
                for block in content_blocks:
                    if not block.strip():
                        continue
                    if '\\begin{figure}' in block:
                        img_match = re.search(r'\\includegraphics.*?\{([^}]+)\}', block)
                        if img_match:
                            img_path = output_dir / img_match.group(1).replace('/', os.sep)
                            if img_path.exists():
                                try:
                                    doc.add_picture(str(img_path), width=Inches(5))
                                except:
                                    pass
                        continue
                    if block.startswith('\\section{') or block.startswith('\\subsection{'):
                        title_text = re.sub(r'\\[a-z]+\{([^}]+)\}', r'\1', block)
                        doc.add_heading(title_text.strip(), level=2 if 'sub' in block else 1)
                    else:
                        for line in block.split('\n'):
                            line = line.strip()
                            if line and not line.startswith('\\'):
                                text = re.sub(r'\\[a-zA-Z]+\{([^}]+)\}', r'\1', line)
                                text = re.sub(r'\\[a-zA-Z]+', '', text)
                                if text.strip():
                                    doc.add_paragraph(text.strip())
                doc.add_page_break()
            
            doc.save(temp_docx)
            print(f"\nWord 文档已保存: {temp_docx}")
            print("提示：请使用 Word 将文件另存为 PDF，或选择 HTML 格式在浏览器中打印为 PDF")
            return temp_docx
        except Exception as e:
            print(f"生成 PDF 失败: {e}，将生成 HTML")
            output_format = 'html'
    
    if output_format == 'docx':
        # 生成 Word 文档
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import re
            
            doc = Document()
            
            # 设置文档标题
            title = doc.add_heading('文档处理结果', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 处理每页内容
            for page_content in all_latex_content:
                # 提取页面标题
                page_match = re.search(r'\\subsection\{第 (\d+) 页\}', page_content)
                if page_match:
                    doc.add_heading(f'第 {page_match.group(1)} 页', level=2)
                    # 移除标题部分
                    page_content = re.sub(r'\\subsection\{第 \d+ 页\}\s*\n\s*', '', page_content)
                
                # 按块处理内容（处理多行的 figure 环境等）
                content_blocks = re.split(r'(\\begin\{figure\}.*?\\end\{figure\}|\\section\{[^}]+\}|\\subsection\{[^}]+\}|\\subsubsection\{[^}]+\})', page_content, flags=re.DOTALL)
                
                for block in content_blocks:
                    if not block.strip():
                        continue
                    
                    # 处理 figure 环境
                    if '\\begin{figure}' in block:
                        # 提取图片路径
                        img_match = re.search(r'\\includegraphics.*?\{([^}]+)\}', block)
                        if img_match:
                            img_rel_path = img_match.group(1).replace('/', os.sep)
                            img_path = output_dir / img_rel_path
                            if img_path.exists():
                                try:
                                    doc.add_picture(str(img_path), width=Inches(5))
                                    # 提取 caption
                                    caption_match = re.search(r'\\caption\{([^}]+)\}', block)
                                    if caption_match:
                                        para = doc.add_paragraph(caption_match.group(1))
                                        para.style = 'Caption'
                                except Exception as e:
                                    print(f"添加图片失败 {img_path}: {e}")
                                    doc.add_paragraph(f"[图片: {img_rel_path}]")
                        continue
                    
                    # 处理标题
                    if block.startswith('\\section{'):
                        title_text = re.sub(r'\\section\{([^}]+)\}', r'\1', block)
                        doc.add_heading(title_text.strip(), level=1)
                    elif block.startswith('\\subsection{'):
                        title_text = re.sub(r'\\subsection\{([^}]+)\}', r'\1', block)
                        doc.add_heading(title_text.strip(), level=2)
                    elif block.startswith('\\subsubsection{'):
                        title_text = re.sub(r'\\subsubsection\{([^}]+)\}', r'\1', block)
                        doc.add_heading(title_text.strip(), level=3)
                    else:
                        # 处理普通文本内容
                        lines = block.split('\n')
                        for line in lines:
                            line = line.strip()
                            if not line:
                                doc.add_paragraph()
                                continue
                            
                            # 跳过 LaTeX 命令
                            if line.startswith('\\') and not line.startswith('\\text'):
                                continue
                            
                            # 处理表格（简化处理）
                            if line.startswith('|') or '\\begin{table}' in line:
                                doc.add_paragraph(line)
                                continue
                            
                            # 普通文本，移除 LaTeX 转义和命令
                            text = line
                            # 移除 LaTeX 命令但保留内容
                            text = re.sub(r'\\textbf\{([^}]+)\}', r'\1', text)  # 粗体
                            text = re.sub(r'\\textit\{([^}]+)\}', r'\1', text)  # 斜体
                            text = re.sub(r'\\emph\{([^}]+)\}', r'\1', text)  # 强调
                            text = re.sub(r'\\text\{([^}]+)\}', r'\1', text)  # 文本
                            # 移除转义字符（注意：替换字符串中的反斜杠需要转义）
                            text = re.sub(r'\\([&%$#_{}])', r'\1', text)
                            text = re.sub(r'\\textbackslash', r'\\', text)  # 修复：使用原始字符串
                            # 移除其他 LaTeX 命令
                            text = re.sub(r'\\[a-zA-Z]+\{([^}]+)\}', r'\1', text)
                            text = re.sub(r'\\[a-zA-Z]+', '', text)
                            
                            if text.strip():
                                doc.add_paragraph(text.strip())
                
                doc.add_page_break()
            
            doc.save(output_path)
            print(f"\nWord 文档已保存: {output_path}")
            return output_path  # 成功生成 Word 文档，直接返回
        except ImportError:
            print("未安装 python-docx，无法生成 Word 文档，将生成 LaTeX 文件")
            output_format = 'tex'
        except Exception as e:
            import traceback
            print(f"生成 Word 文档失败: {e}")
            print(f"错误详情: {traceback.format_exc()}")
            print("将生成 LaTeX 文件作为备选")
            output_format = 'tex'
    
    if output_format == 'tex':
        # 生成完整的 LaTeX 文档
        latex_document = f"""\\documentclass[12pt]{{article}}
\\usepackage[utf8]{{inputenc}}
\\usepackage{{graphicx}}
\\usepackage{{float}}
\\usepackage{{amsmath}}
\\usepackage{{amsfonts}}
\\usepackage{{amssymb}}
\\usepackage{{booktabs}}
\\usepackage{{array}}
\\usepackage{{longtable}}
\\usepackage{{xcolor}}
\\usepackage{{geometry}}
\\geometry{{a4paper, margin=2.5cm}}

\\title{{文档处理结果}}
\\author{{自动生成}}
\\date{{\\today}}

\\begin{{document}}

\\maketitle

{body_content}

\\end{{document}}"""
        
        # 保存 LaTeX 文件
        latex_path = output_path.with_suffix('.tex')
        with open(latex_path, 'w', encoding='utf-8') as f:
            f.write(latex_document)
        print(f"\nLaTeX 文件已保存: {latex_path}")
        
        # 同时保存一个纯内容版本（不含文档结构）
        content_path = output_path.with_suffix('.tex.content')
        with open(content_path, 'w', encoding='utf-8') as f:
            f.write(body_content)
        print(f"LaTeX 内容文件已保存: {content_path}")
        
        return latex_path  # 返回生成的文件路径

